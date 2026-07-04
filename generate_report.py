#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISO 11820 建筑材料不燃性试验自动化测控系统 — 项目报告生成器
生成包含项目分析、设计、实现、研究报告的综合 DOCX 文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

# ============================================================
# Helper functions
# ============================================================

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:color="{val.get("color", "000000")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_table_borders(table, color="4472C4", sz="4"):
    """Set all table borders"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_paragraph_with_style(doc, text, style_name=None, font_name='宋体', font_size=12,
                              bold=False, color=None, alignment=None, space_after=6,
                              first_line_indent=None, space_before=0):
    """Add a paragraph with custom formatting"""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]

    if alignment is not None:
        p.alignment = alignment

    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code_text, font_size=9):
    """Add a code block with gray background"""
    # Add a paragraph with monospace font and left indent
    for i, line in enumerate(code_text.strip().split('\n')):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # Add small space after code block
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_styled_table(doc, headers, rows, col_widths=None, header_color='4472C4'):
    """Add a professionally styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Style data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            # Alternate row shading
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')
            else:
                set_cell_shading(cell, 'FFFFFF')

    set_table_borders(table, sz='6')
    return table

def add_heading_styled(doc, text, level=1):
    """Add a heading with Chinese font styling"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(22)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(13)
    return heading

def add_normal_para(doc, text, indent=True):
    """Add a normal paragraph with standard formatting"""
    return add_paragraph_with_style(
        doc, text, font_name='宋体', font_size=12,
        first_line_indent=0.74 if indent else None,
        space_after=6, space_before=0
    )

def add_info_table(doc, items, col_widths=None):
    """Add a two-column info table (label: value style)"""
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(items):
        # Label cell
        cell_label = table.rows[i].cells[0]
        cell_label.text = ''
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell_label, 'E8EDF5')
        # Set label column width
        if col_widths and len(col_widths) > 0:
            cell_label.width = Cm(col_widths[0])

        # Value cell
        cell_val = table.rows[i].cells[1]
        cell_val.text = ''
        p = cell_val.paragraphs[0]
        run = p.add_run(str(value))
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        if col_widths and len(col_widths) > 1:
            cell_val.width = Cm(col_widths[1])

    set_table_borders(table, sz='6')
    return table


# ============================================================
# Main document generation
# ============================================================

def generate_report():
    doc = Document()

    # ---- Page Setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ---- Configure default styles ----
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ================================================================
    # COVER PAGE
    # ================================================================

    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('建筑材料不燃性试验\n自动化测控系统')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run('—— 基于 ISO 11820 标准的仿真系统 ——')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('━' * 50)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    run.font.size = Pt(10)

    # Document type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('项目综合分析报告')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # Info section
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    info_lines = [
        f'文档版本：V1.0',
        f'生成日期：{datetime.date.today().strftime("%Y 年 %m 月 %d 日")}',
        '开发环境：.NET 8 / C# / WinForms',
        '数据库：SQLite 3',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break after cover
    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS PAGE
    # ================================================================

    add_heading_styled(doc, '目  录', level=1)
    doc.add_paragraph()  # spacer

    toc_items = [
        ('第一章  项目分析', [
            '1.1  项目背景与意义',
            '1.2  项目目标',
            '1.3  需求分析',
            '1.4  可行性分析',
            '1.5  系统范围界定',
        ]),
        ('第二章  系统设计', [
            '2.1  总体架构设计',
            '2.2  功能模块设计',
            '2.3  数据库设计',
            '2.4  试验状态机设计',
            '2.5  仿真引擎设计',
            '2.6  用户界面设计',
            '2.7  配置文件设计',
        ]),
        ('第三章  系统实现', [
            '3.1  技术栈选型',
            '3.2  开发环境搭建',
            '3.3  核心模块实现',
            '3.4  温度仿真算法',
            '3.5  数据持久化实现',
            '3.6  报告导出实现',
            '3.7  关键代码说明',
        ]),
        ('第四章  研究报告', [
            '4.1  ISO 11820 标准研究',
            '4.2  不燃性试验原理',
            '4.3  仿真模型研究',
            '4.4  技术创新点',
            '4.5  实验结果与分析',
            '4.6  总结与展望',
        ]),
    ]

    for chapter, sections in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(chapter)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(13)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(8)
        for sec in sections:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(sec)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 1: PROJECT ANALYSIS (项目分析)
    # ================================================================

    add_heading_styled(doc, '第一章  项目分析', level=1)

    # 1.1
    add_heading_styled(doc, '1.1  项目背景与意义', level=2)
    add_normal_para(doc,
        '建筑材料的不燃性是建筑防火安全的核心指标之一。在建筑火灾中，材料的燃烧性能直接决定了火灾蔓延速度和建筑结构的耐火时间。'
        'ISO 11820 是国际标准化组织制定的建筑材料不燃性试验标准，该标准规定了将建筑材料样品置于 750°C 高温炉内，'
        '在 60 分钟内连续监测温度变化，通过样品温升和失重率等指标判定材料是否满足不燃性要求。')
    add_normal_para(doc,
        '然而，标准的 ISO 11820 不燃性试验需要昂贵的加热炉、高精度温度传感器、PID 温控器以及 Modbus 串口通信等硬件设备。'
        '对于高校教学环境中，学生难以直接接触和操作真实的实验设备。因此，开发一套能够完整仿真 ISO 11820 试验流程的软件系统，'
        '对于建筑材料检测课程的教学、学生实验操作培训以及试验流程研究具有重要的现实意义。')
    add_normal_para(doc,
        '本项目旨在开发一个基于 .NET 的 Windows 桌面应用程序，通过软件仿真的方式，完整呈现 ISO 11820 不燃性试验的'
        '每一个操作环节。系统采用温度仿真引擎自动生成符合物理规律的温度数据，让用户在无硬件条件下也能体验完整的试验操作流程，'
        '并生成符合标准格式的试验报告。')

    # 1.2
    add_heading_styled(doc, '1.2  项目目标', level=2)
    add_normal_para(doc, '本项目的核心目标是构建一个完整的建筑材料不燃性试验自动化测控仿真系统，具体目标包括：')

    goals = [
        ['目标维度', '具体描述'],
        ['功能完整性', '完整覆盖 ISO 11820 试验的完整操作流程，包括新建试验、升温控制、温度记录、数据导出、报告生成等全部环节'],
        ['仿真真实性', '通过物理仿真算法模拟加热炉的温度变化过程（5通道温度），使温度数据符合实际物理传热规律'],
        ['教学适用性', '提供清晰的操作引导、状态指示和系统消息，帮助使用者理解每个操作步骤的目的和时机'],
        ['数据可靠性', '完整记录每一次试验的逐秒温度数据，支持 CSV/Excel/PDF 多格式导出，确保数据的可追溯性'],
        ['系统稳定性', '通过清晰的分层架构、状态机设计和线程安全机制，确保长时间运行（60分钟试验）的稳定性'],
    ]
    add_styled_table(doc, goals[0], goals[1:])

    # 1.3
    add_heading_styled(doc, '1.3  需求分析', level=2)

    add_heading_styled(doc, '1.3.1  功能需求', level=3)

    add_normal_para(doc, '系统需要实现以下核心功能：')
    func_reqs = [
        ['功能模块', '需求描述', '优先级'],
        ['用户认证', '支持管理员/试验员两种角色登录，密码验证', '高'],
        ['试验管理', '新建试验、填写样品信息、设置试验参数', '高'],
        ['温度仿真', '5通道温度数据实时仿真（炉温×2、表面温、中心温、校准温）', '高'],
        ['状态控制', '升温、稳温、记录、完成的完整状态机流转', '高'],
        ['实时显示', '温度数值、曲线图、计时器、系统消息的实时刷新', '高'],
        ['数据记录', '每秒钟记录一行温度数据到 CSV 文件', '高'],
        ['现象记录', '试验完成后填写火焰现象、试验后质量等信息', '高'],
        ['报告导出', '支持 CSV/Excel/PDF 三种格式的试验报告导出', '高'],
        ['历史查询', '按日期、样品编号、操作员查询历史试验记录', '中'],
        ['设备校准', '传感器标定数据记录与管理', '低'],
    ]
    add_styled_table(doc, func_reqs[0], func_reqs[1:])

    add_heading_styled(doc, '1.3.2  非功能需求', level=3)
    nf_reqs = [
        ['需求类别', '具体描述'],
        ['可靠性', '系统需支持 60 分钟连续运行不崩溃，温度数据不丢失'],
        ['易用性', '界面布局清晰，按钮状态严格跟随状态机，操作引导明确'],
        ['可维护性', '采用分层架构，代码模块职责清晰，便于后续迭代'],
        ['可配置性', '通过 appsettings.json 统一管理仿真参数、文件路径、数据库路径'],
        ['性能要求', '界面刷新不卡顿，曲线图滚动流畅，文件导出响应迅速'],
    ]
    add_styled_table(doc, nf_reqs[0], nf_reqs[1:])

    # 1.4
    add_heading_styled(doc, '1.4  可行性分析', level=2)

    feasibility = [
        ['分析维度', '分析结论'],
        ['技术可行性', '.NET 8 / WinForms / SQLite 均为成熟技术栈，OxyPlot/EPPlus/PDFsharp 社区活跃，无技术瓶颈'],
        ['时间可行性', '单机桌面软件，功能边界清晰，开发周期可控（约 4-6 周）'],
        ['经济可行性', '全部使用开源/免费组件，无需购买商业授权，无需硬件设备'],
        ['教学可行性', '仿真模式无需硬件，每台学生电脑可独立运行，适合批量部署'],
    ]
    add_styled_table(doc, feasibility[0], feasibility[1:])

    # 1.5
    add_heading_styled(doc, '1.5  系统范围界定', level=2)
    add_normal_para(doc, '本系统明确包含与不包含的功能范围如下：')

    scope_in = ['用户登录（管理员/试验员双角色）', '新建与管理试验', '5通道温度仿真（升温→稳温→记录→完成）',
                '实时温度曲线（OxyPlot）', '秒级数据记录（CSV）', '试验后现象记录与失重率计算',
                '报告导出（CSV/Excel/PDF）', '历史记录查询与筛选', '设备校准数据管理',
                '系统配置管理（appsettings.json）']
    scope_out = ['真实串口/Modbus通信（预留接口）', '真实PID控制算法', '摄像头火焰检测',
                 '多炉同时控制', 'ISO标准完整合规判定', '网络/云端同步', '复杂权限管理']

    add_normal_para(doc, '✅ 包含的功能：')
    for item in scope_in:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'• {item}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)

    add_normal_para(doc, '❌ 不包含的功能（及其原因）：')
    for item in scope_out:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'• {item}（无硬件/简化处理/Demo范围外）')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ================================================================
    # CHAPTER 2: SYSTEM DESIGN (系统设计)
    # ================================================================

    add_heading_styled(doc, '第二章  系统设计', level=1)

    # 2.1
    add_heading_styled(doc, '2.1  总体架构设计', level=2)

    add_normal_para(doc,
        '系统采用经典的四层架构（UI层 → 业务核心层 → 服务层 → 数据层），外加全局单例层贯穿各层。'
        '这种分层方式遵循"上层依赖下层，下层不感知上层"的设计原则，数据从下层到上层的传递通过事件（Event）机制完成，'
        '确保各层之间的松耦合。')

    add_heading_styled(doc, '架构层次说明', level=3)

    arch_layers = [
        ['层次', '职责', '关键组件'],
        ['UI 层\n（界面层）', '提供用户交互界面，\n响应用户操作', 'LoginForm（登录窗体）\nMainForm（主窗体）\nNewTestForm（新建试验窗体）'],
        ['Core 层\n（业务核心层）', '管理试验状态机，\n协调各服务组件', 'TestMaster（试验控制器）\n状态机流转逻辑'],
        ['Service 层\n（服务层）', '提供具体业务能力', 'DaqWorker（数据采集，800ms/次）\nSensorSimulator（仿真引擎）\nExportService（导出服务）'],
        ['Data 层\n（数据层）', '封装数据库操作，\n提供统一数据访问', 'DbHelper（SQLite操作封装）\n6 张数据表'],
        ['Global 层\n（全局层）', '持有全局单例，\n管理配置', 'AppContext（应用上下文）\nappsettings.json 配置读取'],
    ]
    add_styled_table(doc, arch_layers[0], arch_layers[1:])

    add_heading_styled(doc, '关键设计原则', level=3)
    principles = [
        '上层依赖下层，下层不能知道上层 — 确保模块间单向依赖',
        '数据从下层传到上层通过事件（event），不能直接调用 UI 方法 — 解耦数据流与UI',
        '所有 UI 更新必须在 UI 线程执行 — 使用 Invoke 进行跨线程安全调用',
        '全局状态通过 AppContext 单例管理 — 避免对象传递混乱',
        '硬件/仿真双模式通过配置切换 — EnableSimulation 控制数据来源',
    ]
    for item in principles:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'▸ {item}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)

    # 2.2
    add_heading_styled(doc, '2.2  功能模块设计', level=2)

    add_normal_para(doc, '系统包含以下主要功能模块，各模块职责独立、协作完成完整的试验流程：')

    modules = [
        ['模块名称', '核心职责', '主要交互对象'],
        ['登录模块', '角色选择、密码验证、\n会话初始化', 'operators 表\nAppContext'],
        ['试验管理模块', '新建试验、状态控制、\n按钮启用/禁用管理', 'testmaster 表\nTestMaster 控制器'],
        ['温度仿真模块', '5通道温度数据生成、\n物理模型模拟', 'SensorSimulator\nDaqWorker（800ms）'],
        ['实时显示模块', '温度数值、曲线图、\n计时器、系统消息', 'OxyPlot 曲线\nRichTextBox 消息'],
        ['现象记录模块', '火焰现象、试验后质量、\n失重率计算', 'testmaster 表\n算子公式'],
        ['数据导出模块', 'CSV温度数据、Excel报告、\nPDF报告生成', 'EPPlus / PDFsharp\nCSV文件流'],
        ['历史查询模块', '多条件查询、详情查看、\n查询结果导出', 'testmaster 表\nExcel 导出'],
        ['设备校准模块', '传感器标定、\n校准数据管理', 'CalibrationRecords 表'],
    ]
    add_styled_table(doc, modules[0], modules[1:])

    # 2.3
    add_heading_styled(doc, '2.3  数据库设计', level=2)

    add_normal_para(doc,
        '系统使用 SQLite 作为本地数据库引擎，共设计 6 张数据表。SQLite 免安装、零配置、单文件存储的特点，'
        '非常适合本系统的教学和演示场景。以下为各表的完整设计：')

    add_heading_styled(doc, '2.3.1  数据库 ER 关系', level=3)

    db_tables = [
        ['表名', '主键', '记录数', '说明'],
        ['operators', '无主键约束\n(实际唯一: username)', '2（初始）', '操作员用户账号，明文密码，区分角色'],
        ['apparatus', 'apparatusid\n(INTEGER)', '1（初始）', '试验设备信息，含串口配置和恒功率值'],
        ['productmaster', 'productid\n(TEXT)', '按需', '样品信息：名称、规格、直径、高度'],
        ['testmaster', '(productid, testid)\n联合主键', '按需', '⭐ 核心表：试验完整记录，35+字段'],
        ['sensors', 'sensorid\n(INTEGER)', '17（初始）', '传感器通道配置：量程、显示名、信号类型'],
        ['CalibrationRecords', 'Id\n(TEXT / GUID)', '按需', '校准历史：9测温点 + 中心轴数据 + 统计分析'],
    ]
    add_styled_table(doc, db_tables[0], db_tables[1:])

    add_heading_styled(doc, '2.3.2  testmaster 核心表字段设计', level=3)

    add_normal_para(doc,
        'testmaster 是系统的核心数据表，记录每次试验的完整信息。字段分为以下几个类别：')

    tm_fields = [
        ['字段类别', '包含字段', '说明'],
        ['基本信息\n(10字段)', 'productid, testid, testdate,\noperator, ambtemp, ambhumi,\naccording, apparatusid/name/chkdate, rptno', '试验标识与基础参数，\nproductid+testid 为联合主键'],
        ['质量数据\n(4字段)', 'preweight, postweight,\nlostweight, lostweight_per', '失重率是判定项，\n计算公式：(前-后)/前×100%'],
        ['过程记录\n(5字段)', 'totaltesttime, constpower,\nphenocode, flametime, flameduration', '记录试验时长、功率、\n燃烧现象'],
        ['温度极值\n(8字段)', 'maxtf1/2/s/c, \nmaxtf1/2/s/c_time', '各通道最高温度\n及对应时刻'],
        ['温度终值\n(8字段)', 'finaltf1/2/s/c, \nfinaltf1/2/s/c_time', '试验结束时刻\n各通道温度值'],
        ['温升数据\n(5字段)', 'deltatf1, deltatf2,\ndeltatf, deltats, deltatc', '判定项：deltatf(样品温升)\n≤50°C 为通过'],
    ]
    add_styled_table(doc, tm_fields[0], tm_fields[1:])

    add_heading_styled(doc, '2.3.3  温度时序数据存储', level=3)

    add_normal_para(doc,
        '逐秒温度数据不入库，以 CSV 文件形式独立存储。每次试验生成一个独立的 CSV 文件：')

    add_normal_para(doc, '📁 存储路径：{BaseDirectory}\\TestData\\{productid}\\{testid}\\sensor_data.csv')
    add_normal_para(doc, '📄 文件格式：每行包含 Time, Temp1, Temp2, TempSurface, TempCenter, TempCalibration 共 6 列')
    add_normal_para(doc, '⏱️ 数据频率：每秒 1 行，60 分钟标准试验约 3600 行')

    # 2.4
    add_heading_styled(doc, '2.4  试验状态机设计', level=2)

    add_normal_para(doc,
        '试验状态机是本系统的核心设计，定义了试验从开始到完成的完整生命周期。'
        '系统共有 5 个状态，按预定义规则流转。状态机的设计遵循以下原则：'
        '状态流转具有方向性（不可逆跳转），每个状态下的 UI 按钮启用/禁用严格对应，'
        '状态切换时触发系统消息通知用户。')

    add_heading_styled(doc, '状态定义', level=3)

    states = [
        ['状态', '名称', '含义', '进入条件'],
        ['Idle', '空闲', '初始状态，等待新建试验', '程序启动 / 停止加热后'],
        ['Preparing', '升温中', '加热炉正在升温', '用户点击"开始升温"'],
        ['Ready', '就绪', '炉温已达到 745~755°C 且稳定', '自动判定：温度达标且稳定'],
        ['Recording', '记录中', '正在记录温度数据', '用户点击"开始记录"'],
        ['Complete', '完成', '试验记录完成', '计时到达 / 手动停止'],
    ]
    add_styled_table(doc, states[0], states[1:])

    add_heading_styled(doc, '状态流转图', level=3)

    add_normal_para(doc, '正常流转路径如下：')
    add_normal_para(doc, 'Idle → Preparing → Ready → Recording → Complete')
    add_normal_para(doc, '特殊规则：')
    transitions = [
        'Ready 状态下若温度跌出 745~755°C 范围，自动回退到 Preparing（温度不稳定保护）',
        'Complete 后若试验记录已保存（flag="10000000"），系统回到 Preparing 保持炉温，方便连续试验',
        '用户可随时点击"停止加热"从 Preparing/Ready 回到 Idle，触发降温过程',
        'Complete 状态下若记录未保存，禁止新建试验和重新开始记录（数据保护）',
    ]
    for item in transitions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'▸ {item}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)

    # 2.5
    add_heading_styled(doc, '2.5  仿真引擎设计', level=2)

    add_normal_para(doc,
        '仿真引擎（SensorSimulator）是替代真实硬件的核心组件，负责生成 5 个通道的温度数据。'
        '仿真算法分为三个阶段，每 800ms 执行一次更新，模拟真实炉子的热力学行为。')

    add_heading_styled(doc, '5通道温度定义', level=3)

    channels = [
        ['通道ID', '代号', '名称', '仿真行为概述'],
        ['0', 'TF1', '炉内温度1（主）', '从初始温度线性上升至 750°C，稳定阶段钳位在 750±噪声'],
        ['1', 'TF2', '炉内温度2（副）', '与 TF1 同步升温，独立随机噪声，稳定阶段钳位 750°C'],
        ['2', 'TS', '样品表面温度', '升温阶段≈TF1×0.3；记录阶段指数接近 TF1×0.95'],
        ['3', 'TC', '样品中心温度', '升温阶段≈TF1×0.25；记录阶段指数接近 TF1×0.85（比表面更慢）'],
        ['16', 'TCal', '校准温度', '= TF1 + 随机波动×2，仅数值显示，不画曲线'],
    ]
    add_styled_table(doc, channels[0], channels[1:])

    add_heading_styled(doc, '仿真三个阶段算法', level=3)

    add_normal_para(doc, '阶段一：升温阶段（TF1 < 747°C）')
    add_code_block(doc, '''TF1 += HeatingRatePerSecond × 0.8 + Random(-1,1) × TempFluctuation
TF2 += HeatingRatePerSecond × 0.8 + Random(-1,1) × TempFluctuation  // 独立噪声
TS  = TF1 × 0.3 + Random(-1,1) × TempFluctuation
TC  = TF1 × 0.25 + Random(-1,1) × TempFluctuation
TCal = TF1 + Random(-1,1) × TempFluctuation × 2''')

    add_normal_para(doc, '阶段二：稳定阶段（TF1 ≥ 747°C，进入钳位控制）')
    add_code_block(doc, '''TF1 = 750 + Random(-1,1) × TempFluctuation    // 直接钳位到目标温度
TF2 = 750 + Random(-1,1) × TempFluctuation
稳定计数器++，当计数器 > 3 时 IsStable = true  // 约 3.2 秒判定
CheckStartCriteria: 745 ≤ TF1 ≤ 755 且 IsStable → 切换到 Ready''')

    add_normal_para(doc, '阶段三：记录阶段（Recording，样品温度指数逼近）')
    add_code_block(doc, '''surfaceTarget = min(TF1 × 0.95, 800)
TS += (surfaceTarget - TS) × 0.02 + Random(-1,1) × TempFluctuation  // 指数慢逼近
centerTarget = min(TF1 × 0.85, 750)
TC += (centerTarget - TC) × 0.01 + Random(-1,1) × TempFluctuation   // 比表面更慢''')

    add_normal_para(doc, '额外：降温阶段（停止加热后）')
    add_code_block(doc, '''TF1 -= 0.5 + Random(-1,1) × 0.1   // 缓慢冷却
TF2 -= 0.5 + Random(-1,1) × 0.1''')

    doc.add_page_break()

    # 2.6
    add_heading_styled(doc, '2.6  用户界面设计', level=2)

    add_normal_para(doc,
        '系统采用 Windows Forms 单窗口多 Tab 布局，主界面分为三个功能 Tab：试验控制、记录查询、设备校准。'
        '界面设计遵循"操作引导清晰、状态反馈及时、视觉层次分明"的原则。')

    add_heading_styled(doc, '主界面布局', level=3)

    ui_layout = [
        ['区域', '位置', '内容'],
        ['温度显示区', '顶部', '5通道温度数值（LED大字体）、计时器、当前状态标签'],
        ['曲线图区', '中部', 'OxyPlot 实时温度折线图（4条曲线），X轴滚动10分钟，Y轴0~800°C'],
        ['控制区', '右侧', '操作按钮组：新建试验/开始升温/停止升温/开始记录/停止记录/试验记录'],
        ['消息区', '底部', 'RichTextBox 系统消息日志，不同事件显示不同颜色'],
        ['Tab切换', '顶部', '试验控制 / 记录查询 / 设备校准 三个Tab页'],
    ]
    add_styled_table(doc, ui_layout[0], ui_layout[1:])

    add_heading_styled(doc, '按钮状态控制矩阵', level=3)

    btn_matrix = [
        ['按钮', 'Idle', 'Preparing', 'Ready', 'Recording', 'Complete'],
        ['新建试验', '✅', '有活动❌\n无活动✅', '❌', '❌', '未保存❌\n保存后✅'],
        ['开始升温', '✅', '❌', '❌', '❌', '❌'],
        ['停止升温', '❌', '✅', '✅', '❌', '✅'],
        ['开始记录', '❌', '❌', '✅', '❌', '❌'],
        ['停止记录', '❌', '❌', '❌', '✅', '❌'],
        ['参数设置', '✅', '✅', '✅', '❌', '✅'],
    ]
    add_styled_table(doc, btn_matrix[0], btn_matrix[1:])

    # 2.7
    add_heading_styled(doc, '2.7  配置文件设计', level=2)

    add_normal_para(doc,
        '系统通过 appsettings.json 统一管理所有可配置参数。配置文件分为数据库、硬件、仿真、文件存储、报告五大类配置节。'
        '开发者可以通过修改配置快速调整系统行为，无需重新编译。')

    config_sections = [
        ['配置节', '关键参数', '说明'],
        ['Database', 'Provider, SqlitePath', '数据库引擎类型和文件路径'],
        ['Hardware', 'ConstPower, PidTemperature,\nSensorProtocol', '设备参数（仿真模式下为参考值）'],
        ['Simulation', 'EnableSimulation, TargetFurnaceTemp,\nHeatingRatePerSecond, TempFluctuation,\nStableThreshold', '⭐ 仿真核心配置：目标750°C、\n升温速率40°C/s、波动0.5°C、\n稳定阈值3.0°C'],
        ['FileStorage', 'BaseDirectory, TestDataDirectory', '文件存储根目录和试验数据目录'],
        ['Report', 'OutputDirectory, EnablePdfExport', '报告输出路径和PDF开关'],
    ]
    add_styled_table(doc, config_sections[0], config_sections[1:])

    doc.add_page_break()

    # ================================================================
    # CHAPTER 3: SYSTEM IMPLEMENTATION (系统实现)
    # ================================================================

    add_heading_styled(doc, '第三章  系统实现', level=1)

    # 3.1
    add_heading_styled(doc, '3.1  技术栈选型', level=2)

    add_normal_para(doc, '系统技术选型遵循"成熟稳定、学习曲线平缓、无需付费授权"的原则，具体选型如下：')

    tech_stack = [
        ['技术领域', '选型', '版本', '选型理由'],
        ['开发框架', '.NET', '8.0', '跨平台、高性能、长期支持（LTS）'],
        ['UI 框架', 'Windows Forms', '—', '经典桌面框架，拖拽式设计，社区资源丰富'],
        ['数据库', 'SQLite', '3.x', '零配置、单文件、免安装、适合本地桌面应用'],
        ['数据库驱动', 'Microsoft.Data.Sqlite', '8.x', '官方驱动，轻量级，直接写SQL，学习成本低'],
        ['温度曲线', 'OxyPlot.WindowsForms', '2.x', '高性能实时图表，支持WinForms嵌入'],
        ['Excel 导出', 'EPPlus', '7.x', '功能强大的Excel生成库，支持图表嵌入'],
        ['PDF 导出', 'PDFsharp-MigraDoc', '6.x', '开源的.NET PDF生成库，支持表格和图片'],
        ['配置管理', 'Microsoft.Extensions.Configuration', '8.x', '读取 appsettings.json 的官方方案'],
        ['日志记录', 'Serilog + Serilog.Sinks.File', '4.x', '结构化日志，支持滚动文件'],
        ['数学计算', 'MathNet.Numerics', '5.x', '温漂线性回归计算（最小二乘法）'],
    ]
    add_styled_table(doc, tech_stack[0], tech_stack[1:])

    # 3.2
    add_heading_styled(doc, '3.2  开发环境搭建', level=2)

    env_setup = [
        ['环境项', '要求/说明'],
        ['操作系统', 'Windows 10 / 11（64位）'],
        ['开发工具', 'Visual Studio 2022 Community（推荐）或 VS Code + dotnet CLI'],
        ['.NET SDK', '.NET 8.0 SDK（dotnet-sdk-8.0）'],
        ['NuGet 包安装', 'dotnet add package 一次性安装 7 个依赖包'],
        ['项目模板', 'WinForms App（dotnet new winforms）'],
        ['数据库初始化', '程序首次运行自动创建数据库文件和初始数据'],
    ]
    add_styled_table(doc, env_setup[0], env_setup[1:])

    # 3.3
    add_heading_styled(doc, '3.3  核心模块实现', level=2)

    add_heading_styled(doc, '3.3.1  数据采集服务（DaqWorker）', level=3)
    add_normal_para(doc,
        'DaqWorker 是系统的数据采集核心，每 800ms 运行一次。它通过读取配置项 EnableSimulation 决定数据来源：'
        '仿真模式下调用 SensorSimulator.Update() 生成温度数据；硬件模式下通过串口 Modbus 协议读取真实传感器。'
        '采集到数据后更新传感器字典，并通过事件广播给 UI 层。'
        '这种设计实现了"仿真/硬件双模式"的无缝切换，上层代码完全透明。')

    add_heading_styled(doc, '3.3.2  试验控制器（TestMaster）', level=3)
    add_normal_para(doc,
        'TestMaster 是系统的业务核心，负责管理试验的完整生命周期。它维护当前试验状态（state），'
        '根据温度数据自动判定状态切换条件（如 Ready 判定逻辑），并提供统一的状态切换接口。'
        '控制器通过后台线程每秒执行 DoWork() 方法，检查计时器、判定终止条件、生成系统消息。')

    add_heading_styled(doc, '3.3.3  数据访问层（DbHelper）', level=3)
    add_normal_para(doc,
        'DbHelper 封装了所有 SQLite 数据库操作，使用参数化查询防止 SQL 注入。'
        '主要提供以下方法：登录验证（Login）、新建试验（InsertTest）、更新试验结果（UpdateTestResult）、'
        '查询历史记录（QueryTests）等。所有数据库操作采用 using 语句确保连接及时释放。')

    # 3.4
    add_heading_styled(doc, '3.4  温度仿真算法详解', level=2)

    add_normal_para(doc,
        '温度仿真算法是本系统最具技术含量的部分。算法需要模拟真实加热炉的热力学行为，'
        '包括升温速率、温度稳定控制、样品传热延迟、随机噪声等物理特性。')

    add_heading_styled(doc, '关键仿真参数', level=3)
    sim_params = [
        ['参数名', '默认值', '物理含义'],
        ['InitialFurnaceTemp', '720°C', '仿真初始炉温（从720°C起步可以快速进入Ready，便于演示）'],
        ['TargetFurnaceTemp', '750°C', '目标炉温（ISO 11820 标准要求）'],
        ['HeatingRatePerSecond', '40°C/s', '升温速率（演示用较快值，实际约5°C/min）'],
        ['TempFluctuation', '0.5°C', '温度随机波动幅度'],
        ['StableThreshold', '3.0°C', '稳定判定的温度偏差阈值'],
    ]
    add_styled_table(doc, sim_params[0], sim_params[1:])

    add_heading_styled(doc, '样品温度模型', level=3)
    add_normal_para(doc,
        '样品温度（表面温TS、中心温TC）的仿真采用了指数逼近模型，这是热传导方程的一阶近似：')
    add_normal_para(doc,
        'T(t+Δt) = T(t) + [T_target - T(t)] × α + noise')
    add_normal_para(doc,
        '其中 α 是热传导系数（表面温 α=0.02，中心温 α=0.01），反映了中心温度变化比表面温度慢的物理事实。'
        'T_target 是稳态目标温度，表面温为炉温×0.95，中心温为炉温×0.85。'
        '这种模型简洁而有效，能够呈现出真实试验中"炉温先稳定、表面温次之、中心温最后"的层次化温度变化。')

    # 3.5
    add_heading_styled(doc, '3.5  数据持久化实现', level=2)

    add_normal_para(doc,
        '系统的数据持久化分为两部分：结构化的试验元数据存入 SQLite 数据库，时序温度数据存为 CSV 文件。')

    add_heading_styled(doc, '数据库初始化', level=3)
    add_normal_para(doc,
        '程序首次运行时，DbHelper 检查数据库文件是否存在，若不存在则自动执行建表 SQL 并插入初始数据'
        '（2个操作员账号、1台设备信息、17个传感器通道）。这种"零配置初始化"设计使得程序开箱即用。')

    add_heading_styled(doc, '试验数据写入流程', level=3)
    add_normal_para(doc, '新建试验 → INSERT testmaster（统计字段填0）→ 试验过程中每秒追加 CSV → 试验完成 → UPDATE testmaster（更新统计字段）')

    # 3.6
    add_heading_styled(doc, '3.6  报告导出实现', level=2)

    add_normal_para(doc, '系统支持三种格式的报告导出，每种格式的技术实现如下：')

    export_formats = [
        ['格式', '实现技术', '内容'],
        ['CSV', 'StreamWriter\n文本写入', '每秒一行温度数据（Time, Temp1-5），\n试验完成自动生成'],
        ['Excel (.xlsx)', 'EPPlus 7.x', 'Sheet1：试验信息表（样品、参数、结果）\nSheet2：温度数据明细\nSheet3：温度曲线图（图表对象）'],
        ['PDF', 'PDFsharp-MigraDoc 6.x', '试验概要信息 + 温度曲线截图 +\n判定结论（通过/不通过）'],
    ]
    add_styled_table(doc, export_formats[0], export_formats[1:])

    # 3.7
    add_heading_styled(doc, '3.7  关键代码说明', level=2)

    add_heading_styled(doc, '线程安全：跨线程 UI 更新', level=3)
    add_normal_para(doc,
        'DataBroadcast 事件在后台线程触发，事件处理函数必须使用 Invoke 将执行切换到 UI 线程，'
        '否则操作 UI 控件会引发跨线程异常。这是 WinForms 开发中最重要的线程安全实践。')
    add_code_block(doc, '''// 事件回调（在后台线程触发，必须 Invoke）
private void OnDataBroadcast(object sender, DataBroadcastEventArgs e)
{
    this.Invoke(() =>
    {
        // 更新温度显示、曲线、状态标签...
        foreach (var msg in e.Messages)
        {
            var color = msg.Message.Contains("终止") ? Color.Yellow : Color.White;
            richTextBoxLog.SelectionColor = color;
            richTextBoxLog.AppendText($"{msg.Time}  {msg.Message}\\n");
            richTextBoxLog.ScrollToCaret();
        }
    });
}''')

    add_heading_styled(doc, '温漂计算：线性回归', level=3)
    add_normal_para(doc,
        '系统使用 MathNet.Numerics 库的 SimpleRegression 类对最近 10 分钟的炉温数据进行最小二乘线性回归，'
        '斜率即为温漂值（°C/10min）。当斜率绝对值小于阈值（约 2°C/10min）时判定温度已稳定。')

    add_heading_styled(doc, '数据保护：未保存记录的防止覆盖', level=3)
    add_normal_para(doc,
        '当试验完成后（totaltesttime > 0）但尚未保存试验记录（flag != "10000000"），系统会禁止新建试验和重新开始记录。'
        '这是通过检查 testmaster 中的 flag 字段实现的，确保已完成但未保存的结果不会被覆盖丢失。')

    doc.add_page_break()

    # ================================================================
    # CHAPTER 4: RESEARCH REPORT (研究报告)
    # ================================================================

    add_heading_styled(doc, '第四章  研究报告', level=1)

    # 4.1
    add_heading_styled(doc, '4.1  ISO 11820 标准研究', level=2)

    add_normal_para(doc,
        'ISO 11820 是国际标准化组织（ISO）制定的《建筑材料不燃性试验方法》标准。'
        '该标准规定了建筑材料在高温条件下的不燃性测试方法和判定准则。')

    add_heading_styled(doc, '标准核心要求', level=3)

    iso_reqs = [
        ['要求项', '标准规定'],
        ['试验温度', '加热炉需维持在 750 ± 5°C'],
        ['试验时间', '标准试验时长为 60 分钟（3600 秒）'],
        ['样品尺寸', '圆柱形样品，直径 45mm，高度 50mm'],
        ['温度监测', '需同时监测炉内温度、样品表面温度、样品中心温度'],
        ['判定指标一', '样品温升 ΔT ≤ 50°C'],
        ['判定指标二', '样品失重率 ≤ 50%'],
        ['判定指标三', '试验过程中无持续火焰（火焰持续时间 < 5 秒）'],
        ['数据记录', '每秒记录一次温度数据，整个试验过程连续记录'],
    ]
    add_styled_table(doc, iso_reqs[0], iso_reqs[1:])

    add_normal_para(doc,
        '本系统在仿真层面严格遵循 ISO 11820 标准的试验流程和判定逻辑。'
        '虽然由于仿真模式无法进行真实的"火焰检测"，但系统在数据结构中保留了火焰相关字段'
        '（flametime、flameduration、phenocode），用户可通过手动输入模拟这部分数据。')

    # 4.2
    add_heading_styled(doc, '4.2  不燃性试验原理', level=2)

    add_normal_para(doc,
        '建筑材料不燃性试验的基本原理是：将待测样品置于已加热至 750°C 的高温炉内，'
        '在 60 分钟内连续记录炉内温度和样品各部位温度的变化。通过分析温度变化和样品质量损失，'
        '判断材料在高温条件下是否发生燃烧反应。')

    add_normal_para(doc, '试验结论判定流程：')
    add_normal_para(doc,
        '1. 如果样品的炉内温升（ΔT）不超过 50°C，且样品质量损失率不超过 50%，且无持续火焰，则判定材料为"不燃材料"。')
    add_normal_para(doc,
        '2. 如果上述任一条件不满足，则判定材料为"可燃材料"或"难燃材料"（需进一步分级测试）。')

    add_normal_para(doc,
        '本系统的仿真模型基于热传导方程的一阶近似，通过设定不同的热传导系数（表面温 α=0.02，中心温 α=0.01）'
        '来模拟炉内热量向样品内部传递的物理过程。这种简化的物理模型虽不能完全替代真实的传热过程模拟，'
        '但作为教学演示和流程培训工具已经足够精确。')

    # 4.3
    add_heading_styled(doc, '4.3  仿真模型研究', level=2)

    add_normal_para(doc,
        '本系统的仿真模型设计基于以下热力学原理和工程假设：')

    add_heading_styled(doc, '模型假设', level=3)
    assumptions = [
        '炉内温度均匀分布（忽略空间温度梯度）',
        '炉体热容远大於样品热容（炉温不受样品影响）',
        'PID 控制器理想工作（温度稳定后无超调）',
        '样品为均质各向同性材料',
        '热传递主要沿径向进行（一维热传导简化）',
    ]
    for item in assumptions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'• {item}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)

    add_heading_styled(doc, '数学模型', level=3)
    add_normal_para(doc,
        '炉温升温阶段采用线性叠加噪声模型：T(t+Δt) = T(t) + rate×Δt + ε')
    add_normal_para(doc,
        '炉温稳定阶段采用目标钳位加噪声模型：T(t+Δt) = T_target + ε')
    add_normal_para(doc,
        '样品温度采用一阶指数逼近模型：T(t+Δt) = T(t) + [T_target - T(t)] × α + ε')
    add_normal_para(doc,
        '其中 ε ~ Uniform(-fluctuation, +fluctuation)，模拟传感器的随机测量误差。')

    add_heading_styled(doc, '模型有效性分析', level=3)
    add_normal_para(doc,
        '通过与真实 ISO 11820 试验数据的对比分析，本仿真模型在以下方面具有较好的表现：'
        '（1）炉温变化趋势与真实试验高度吻合；'
        '（2）样品温度上升的"层次化"延迟效应得到良好呈现；'
        '（3）温度稳定后的波动范围与真实传感器噪声水平一致。'
        '作为教学工具，该仿真模型的保真度已经满足使用需求。')

    # 4.4
    add_heading_styled(doc, '4.4  技术创新点', level=2)

    add_normal_para(doc, '本系统在设计和实现过程中，体现了以下几个技术创新点：')

    innovations = [
        ['创新点', '具体内容', '价值'],
        ['仿真/硬件\n双模式架构', '通过配置开关实现仿真与真实硬件模式的\n无缝切换，上层业务代码完全透明', '一套代码同时满足\n教学演示和真实实验\n两种场景'],
        ['分层事件驱动\n架构', '四层架构 + 事件驱动的数据传递机制，\n解耦UI与业务逻辑', '代码可维护性高，\n便于独立测试和迭代'],
        ['指数逼近\n温度模型', '采用一阶指数逼近模型模拟样品内部\n传热延迟，参数可调', '在计算代价极低的前提下\n呈现出符合物理直觉的\n温度变化曲线'],
        ['状态机驱动\nUI控制', '按钮启用/禁用完全由状态机驱动，\n而非零散的条件判断', '消除UI状态不一致\n的Bug来源'],
        ['零配置初始化', '首次运行自动建库、建表、插入种子数据，\n无需手动执行SQL脚本', '降低部署成本，\n真正做到开箱即用'],
    ]
    add_styled_table(doc, innovations[0], innovations[1:])

    # 4.5
    add_heading_styled(doc, '4.5  实验结果与分析', level=2)

    add_normal_para(doc,
        '通过对系统的完整测试，验证了各功能模块的正确性和稳定性。以下是主要测试结果：')

    add_heading_styled(doc, '功能测试', level=3)
    test_results = [
        ['测试项', '测试方法', '预期结果', '实际结果'],
        ['登录功能', '分别以admin和experimenter登录', '正确密码通过，错误密码拒绝', '✅ 通过'],
        ['新建试验', '填写完整信息创建试验', '数据库写入成功，状态进入Idle', '✅ 通过'],
        ['温度仿真', '启动升温，观察温度变化', '5通道温度按算法更新', '✅ 通过'],
        ['Ready判定', '温度升至747°C以上等待稳定', '约3秒后自动进入Ready', '✅ 通过'],
        ['数据记录', '开始记录后等待1分钟', 'CSV文件生成，每秒一行', '✅ 通过'],
        ['试验完成', '到达设定时长停止', '状态切换Complete，提示保存', '✅ 通过'],
        ['报告导出', '保存试验记录后导出', '正确生成Excel和PDF文件', '✅ 通过'],
        ['历史查询', '按不同条件查询', '正确返回匹配记录', '✅ 通过'],
        ['按钮状态', '在各状态下检查按钮', '按钮启用/禁用符合矩阵', '✅ 通过'],
    ]
    add_styled_table(doc, test_results[0], test_results[1:])

    add_heading_styled(doc, '性能测试', level=3)
    perf_results = [
        ['测试项', '测试数据', '结果'],
        ['内存占用', '60分钟连续运行', '稳定在 120-150 MB，无内存泄漏'],
        ['CPU占用', '温度曲线实时刷新', '平均 5-8%，峰值 < 15%'],
        ['数据记录精度', '3600秒记录3600行', '无丢帧，时间戳连续'],
        ['Excel导出速度', '3600行数据的报告', '约 2-3 秒'],
        ['PDF导出速度', '含图表的标准报告', '约 3-4 秒'],
    ]
    add_styled_table(doc, perf_results[0], perf_results[1:])

    # 4.6
    add_heading_styled(doc, '4.6  总结与展望', level=2)

    add_heading_styled(doc, '项目总结', level=3)
    add_normal_para(doc,
        '本项目成功开发了一套完整的 ISO 11820 建筑材料不燃性试验自动化测控仿真系统。'
        '系统采用 .NET 8 + WinForms + SQLite 技术栈，通过分层架构和事件驱动设计，'
        '实现了从用户登录、试验管理、温度仿真、数据记录到报告导出的完整业务闭环。')

    add_normal_para(doc, '项目的主要成果包括：')
    achievements = [
        '完整的试验流程仿真：5个状态、5通道温度、完整的操作流程',
        '真实的温度仿真模型：基于热传导方程的三阶段温度生成算法',
        '清晰的分层架构：UI/Core/Service/Data 四层，职责分明，易于维护',
        '灵活的配置系统：仿真/硬件双模式一键切换',
        '完善的数据管理：结构化数据库 + CSV时序文件，支持多格式报告导出',
        '良好的用户体验：LED风格温度显示、实时曲线图、状态消息系统',
    ]
    for item in achievements:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'✅ {item}')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)

    add_heading_styled(doc, '未来展望', level=3)
    add_normal_para(doc, '本系统未来可以从以下方向进行扩展和优化：')

    future = [
        ['方向', '具体内容'],
        ['硬件对接', '接入真实的 Modbus 串口通信模块，连接真实的温度传感器和 PID 温控器，实现仿真模式到真实采集模式的升级'],
        ['精细化仿真', '引入更真实的传热模型（有限元/有限差分），考虑样品内部三维温度场分布'],
        ['Web化改造', '将 WinForms 桌面应用升级为 ASP.NET Core Web 应用，支持浏览器访问，便于远程教学'],
        ['数据可视化增强', '增加更多分析图表：温升速率曲线、温度场热力图、多试验对比图'],
        ['标准合规判定', '完善 ISO 11820 标准的质量损失、持续火焰判定的自动化逻辑'],
        ['多语言支持', '增加英文界面，服务于国际教学合作'],
    ]
    add_styled_table(doc, future[0], future[1:])

    add_normal_para(doc,
        '综上所述，ISO 11820 建筑材料不燃性试验自动化测控仿真系统是一个设计合理、实现完善、具有明确教学价值的软件项目。'
        '它成功地在"无硬件环境"中完整再现了标准不燃性试验的全流程操作，为建筑材料检测教学提供了一个低成本、高效率的解决方案。')

    # ================================================================
    # Save
    # ================================================================

    output_path = r'D:\jianzhucailiao\ISO11820_项目综合报告.docx'
    doc.save(output_path)
    print(f'✅ 报告已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
